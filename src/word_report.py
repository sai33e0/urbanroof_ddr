from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt


class WordReportGenerator:

    def _add_bullet_list(self, document, items):
        if not items:
            document.add_paragraph("Not Available")
            return

        for item in items:
            document.add_paragraph(str(item), style="List Bullet")

    def _add_image_cell(self, cell, image_paths):

        cell.text = ""

        if not image_paths:
            cell.text = "Image Not Available"
            return

        for image_path in image_paths:
            if not image_path or image_path == "Image Not Available" or not Path(image_path).exists():
                paragraph = cell.add_paragraph("Image Not Available")
                continue

            paragraph = cell.add_paragraph()
            run = paragraph.add_run()
            run.add_picture(str(image_path), width=Inches(1.6))

    def generate(self, ddr_json, output_path):

        doc = Document()

        doc.add_heading(
            "Detailed Diagnostic Report",
            level=1
        )

        doc.add_paragraph(
            f"Generated : {datetime.now()}"
        )

        doc.add_heading("1. Property Issue Summary", level=2)

        doc.add_paragraph(
            ddr_json.get("property_issue_summary", "Not Available")
        )

        doc.add_heading("2. Area-wise Observations", level=2)

        areas = ddr_json.get("area_wise_observations")
        if areas is None:
            areas = ddr_json.get("areas", [])

        for area in areas:

            if isinstance(area, str):
                doc.add_paragraph(area)
                continue

            area_name = area.get("area_name") or area.get("area") or "Unknown Area"

            doc.add_heading(area_name, level=3)

            observations = area.get("observations", [])

            if isinstance(observations, str):
                doc.add_paragraph(observations)
                continue

            for obs in observations:
                table = doc.add_table(rows=1, cols=9)
                table.style = "Table Grid"

                headers = [
                    "Page",
                    "Issue",
                    "Description",
                    "Severity Assessment",
                    "Severity Reason",
                    "Probable Root Cause",
                    "Recommended Actions",
                    "Inspection Image",
                    "Thermal Image",
                ]

                for cell, header in zip(table.rows[0].cells, headers):
                    cell.text = header

                row = table.add_row().cells

                if isinstance(obs, str):
                    row[0].text = "Not Available"
                    row[1].text = obs
                    row[2].text = "Not Available"
                    row[3].text = "Not Available"
                    row[4].text = "Not Available"
                    row[5].text = "Not Available"
                    row[6].text = "Not Available"
                    row[7].text = "Image Not Available"
                    row[8].text = "Image Not Available"
                else:
                    row[0].text = str(obs.get("page") or "Not Available")
                    row[1].text = str(obs.get("issue") or obs.get("observation") or "Observation")
                    row[2].text = str(obs.get("description") or obs.get("details") or "Not Available")
                    row[3].text = str(obs.get("severity") or "Not Available")
                    row[4].text = str(obs.get("severity_reason") or "Not Available")
                    row[5].text = str(obs.get("root_cause") or "Not Available")
                    row[6].text = "; ".join(obs.get("recommended_actions", [])) if isinstance(obs.get("recommended_actions"), list) else str(obs.get("recommended_actions") or "Not Available")

                    inspection_images = obs.get("inspection_images", [])
                    thermal_images = obs.get("thermal_images", [])

                    self._add_image_cell(row[7], inspection_images)
                    self._add_image_cell(row[8], thermal_images)

                doc.add_paragraph("")

        doc.add_heading("3. Probable Root Cause", level=2)
        self._add_bullet_list(doc, ddr_json.get("probable_root_cause", []))

        doc.add_heading("4. Severity Assessment", level=2)
        self._add_bullet_list(doc, ddr_json.get("severity_assessment", []))

        doc.add_heading("5. Recommended Actions", level=2)
        self._add_bullet_list(doc, ddr_json.get("recommended_actions", []))

        doc.add_heading("6. Additional Notes", level=2)
        self._add_bullet_list(doc, ddr_json.get("additional_notes", []))

        doc.add_heading("7. Missing / Unclear Information", level=2)
        self._add_bullet_list(doc, ddr_json.get("missing_information", []))

        doc.save(output_path)

        return output_path